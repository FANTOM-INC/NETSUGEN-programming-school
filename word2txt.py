# -*- coding: utf-8 -*-
import os
from pathlib import Path
from docx import Document

path = Path.cwd() / "sample_file/sample.docx"
document = Document(path)
text = ''

# 本文中のテキスト抽出
for i, p in enumerate(document.paragraphs):
    text += p.text

# 表中のテキスト抽出
for i, t in enumerate(document.tables):
    for j, r in enumerate(t.rows):
        for k, c in enumerate(r.cells):
            text += c.text

print(text)
