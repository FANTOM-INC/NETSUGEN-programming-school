from PIL import Image
import sys
import pyocr
import pyocr.builders
import openpyxl
import pdf2image
from docx import Document

import os
from pathlib import Path

# PDF用
os.environ["TESSDATA_PREFIX"] = (Path.cwd() / "setting").as_posix()
tools = pyocr.get_available_tools()
if len(tools) == 0:
    print("No OCR tool found")
    sys.exit(1)
tool = tools[0]


def grep(path, text):
    file_list = []
    files = path.glob("*")
    for file in files:
        if file.is_file():

            # pdfの場合はOCRでテキスト化して検索
            if ".pdf" == file.suffix:
                try:
                    content = ""
                    images = pdf2image.convert_from_path(
                        file.as_posix(), dpi=200, fmt='jpg')
                    lang = 'jpn'
                    for image in images:
                        txt = tool.image_to_string(
                            image,
                            lang=lang,
                            builder=pyocr.builders.TextBuilder()
                        )
                        content += txt
                    content = content.replace("\n", "").replace(
                        " ", "").replace("　", "")
                    if text in content:
                        file_list.append(file.as_posix())
                except:
                    pass

            elif ".docx" == file.suffix:
                try:
                    document = Document(file.as_posix())
                    content = ""

                    # 本文中のテキスト抽出
                    for i, p in enumerate(document.paragraphs):
                        content += p.text

                    # 表中のテキスト抽出
                    for i, t in enumerate(document.tables):
                        for j, r in enumerate(t.rows):
                            for k, c in enumerate(r.cells):
                                content += c.text

                    if text in content:
                        file_list.append(file.as_posix())
                except:
                    pass

            elif ".xlsx" == file.suffix:
                wb = openpyxl.load_workbook(file.as_posix())
                sheets = wb.sheetnames
                for sheet_name in sheets:
                    sheet = wb[sheet_name]
                    for r in sheet.iter_rows():
                        for c in r:
                            if "He" in str(c.value):
                                file_list.append(file.as_posix())
                                break

            else:
                try:
                    with open(file, encoding="UTF-8") as f:
                        if text in f.read():
                            file_list.append(file.as_posix())
                except UnicodeDecodeError:
                    pass
                except Exception as e:
                    print(e)
    return file_list


if __name__ == "__main__":
    text = "Hello"
    path = Path(".")
    # このファイルがあるディレクトリで「シート」という文字列を含むファイルを探す
    print(grep(path, text))
